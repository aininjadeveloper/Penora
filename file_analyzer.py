"""
Enhanced File Analysis Service for Penora
Analyzes uploaded files and determines page count for credit calculation
"""

import logging
import os
import tempfile
from io import BytesIO
from docx import Document
import PyPDF2
import zipfile
import json
import xml.etree.ElementTree as ET

class FileAnalyzer:
    """Analyze uploaded files and extract content with page counting"""
    
    SUPPORTED_FORMATS = {
        'txt': {'name': 'Text File', 'mime': ['text/plain']},
        'docx': {'name': 'Word Document', 'mime': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']},
        'doc': {'name': 'Word Document (Legacy)', 'mime': ['application/msword']},
        'pdf': {'name': 'PDF Document', 'mime': ['application/pdf']},
        'rtf': {'name': 'Rich Text Format', 'mime': ['application/rtf']},
        'odt': {'name': 'OpenDocument Text', 'mime': ['application/vnd.oasis.opendocument.text']},
        'html': {'name': 'HTML Document', 'mime': ['text/html']},
        'md': {'name': 'Markdown File', 'mime': ['text/markdown', 'text/x-markdown']}
    }
    
    WORDS_PER_PAGE = 250  # Standard page estimation
    
    @staticmethod
    def analyze_file(file_stream, filename, file_size):
        """
        Analyze uploaded file and return content info with page count
        Returns: dict with success, content, pages, word_count, file_type, etc.
        """
        try:
            # Get file extension
            file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
            
            if file_ext not in FileAnalyzer.SUPPORTED_FORMATS:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_ext}',
                    'supported_formats': list(FileAnalyzer.SUPPORTED_FORMATS.keys())
                }
            
            # Reset stream position
            file_stream.seek(0)
            
            # Extract content based on file type
            if file_ext == 'txt':
                result = FileAnalyzer._analyze_text_file(file_stream, filename, file_size)
            elif file_ext == 'docx':
                result = FileAnalyzer._analyze_docx_file(file_stream, filename, file_size)
            elif file_ext == 'pdf':
                result = FileAnalyzer._analyze_pdf_file(file_stream, filename, file_size)
            elif file_ext in ['html', 'htm']:
                result = FileAnalyzer._analyze_html_file(file_stream, filename, file_size)
            elif file_ext == 'md':
                result = FileAnalyzer._analyze_markdown_file(file_stream, filename, file_size)
            elif file_ext == 'rtf':
                result = FileAnalyzer._analyze_rtf_file(file_stream, filename, file_size)
            else:
                # Fallback to text processing
                result = FileAnalyzer._analyze_text_file(file_stream, filename, file_size)
            
            if result.get('success'):
                # Calculate Ku coins needed (1 Ku coin per page)
                pages = result.get('pages', 1)
                result['ku_coins_needed'] = max(1, pages)  # Minimum 1 Ku coin
                result['credits_needed'] = result['ku_coins_needed']  # Backward compatibility
                result['file_format'] = FileAnalyzer.SUPPORTED_FORMATS[file_ext]['name']
            
            return result
            
        except Exception as e:
            logging.error(f"File analysis error: {e}")
            return {
                'success': False,
                'error': f'Failed to analyze file: {str(e)}',
                'pages': 0,
                'ku_coins_needed': 0,
                'credits_needed': 0  # Backward compatibility
            }
    
    @staticmethod
    def _analyze_text_file(file_stream, filename, file_size):
        """Analyze plain text file"""
        try:
            content = file_stream.read().decode('utf-8')
            word_count = len(content.split())
            pages = max(1, word_count // FileAnalyzer.WORDS_PER_PAGE)
            
            return {
                'success': True,
                'content': content,
                'word_count': word_count,
                'pages': pages,
                'file_type': 'text',
                'file_size': file_size,
                'filename': filename
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            file_stream.seek(0)
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    file_stream.seek(0)
                    content = file_stream.read().decode(encoding)
                    word_count = len(content.split())
                    pages = max(1, word_count // FileAnalyzer.WORDS_PER_PAGE)
                    
                    return {
                        'success': True,
                        'content': content,
                        'word_count': word_count,
                        'pages': pages,
                        'file_type': 'text',
                        'file_size': file_size,
                        'filename': filename,
                        'encoding': encoding
                    }
                except:
                    continue
            
            return {
                'success': False,
                'error': 'Unable to decode text file with supported encodings',
                'pages': 0
            }
    
    @staticmethod
    def _analyze_docx_file(file_stream, filename, file_size):
        """Analyze Word DOCX file"""
        try:
            # Save to temporary file for python-docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                file_stream.seek(0)
                temp_file.write(file_stream.read())
                temp_file_path = temp_file.name
            
            try:
                # Read document
                doc = Document(temp_file_path)
                
                # Extract text from paragraphs
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)
                
                content = '\n'.join(full_text)
                word_count = len(content.split())
                
                # For DOCX, we can also use the actual page count if available
                pages = len(doc.sections)  # Rough estimate
                if pages == 0:
                    pages = max(1, word_count // FileAnalyzer.WORDS_PER_PAGE)
                
                return {
                    'success': True,
                    'content': content,
                    'word_count': word_count,
                    'pages': pages,
                    'file_type': 'docx',
                    'file_size': file_size,
                    'filename': filename,
                    'sections': len(doc.sections),
                    'paragraphs': len(doc.paragraphs)
                }
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                    
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to process DOCX file: {str(e)}',
                'pages': 0
            }
    
    @staticmethod
    def _analyze_pdf_file(file_stream, filename, file_size):
        """Analyze PDF file"""
        try:
            file_stream.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_stream)
            
            # Get actual page count
            page_count = len(pdf_reader.pages)
            
            # Extract text from all pages
            text_content = []
            for page in pdf_reader.pages:
                try:
                    text_content.append(page.extract_text())
                except:
                    continue
            
            content = '\n'.join(text_content)
            word_count = len(content.split())
            
            return {
                'success': True,
                'content': content,
                'word_count': word_count,
                'pages': page_count,  # Use actual PDF page count
                'file_type': 'pdf',
                'file_size': file_size,
                'filename': filename,
                'pdf_pages': page_count
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to process PDF file: {str(e)}',
                'pages': 0
            }
    
    @staticmethod
    def _analyze_html_file(file_stream, filename, file_size):
        """Analyze HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            file_stream.seek(0)
            content = file_stream.read().decode('utf-8')
            
            # Parse HTML and extract text
            soup = BeautifulSoup(content, 'html.parser')
            text_content = soup.get_text()
            
            # Clean up text
            lines = [line.strip() for line in text_content.split('\n') if line.strip()]
            clean_content = '\n'.join(lines)
            
            word_count = len(clean_content.split())
            pages = max(1, word_count // FileAnalyzer.WORDS_PER_PAGE)
            
            return {
                'success': True,
                'content': clean_content,
                'word_count': word_count,
                'pages': pages,
                'file_type': 'html',
                'file_size': file_size,
                'filename': filename
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to process HTML file: {str(e)}',
                'pages': 0
            }
    
    @staticmethod
    def _analyze_markdown_file(file_stream, filename, file_size):
        """Analyze Markdown file"""
        try:
            file_stream.seek(0)
            content = file_stream.read().decode('utf-8')
            
            # Remove markdown formatting for word count
            import re
            clean_content = re.sub(r'[#*`_\[\]()]', '', content)
            clean_content = re.sub(r'\n+', '\n', clean_content)
            
            word_count = len(clean_content.split())
            pages = max(1, word_count // FileAnalyzer.WORDS_PER_PAGE)
            
            return {
                'success': True,
                'content': content,  # Keep original markdown
                'word_count': word_count,
                'pages': pages,
                'file_type': 'markdown',
                'file_size': file_size,
                'filename': filename
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to process Markdown file: {str(e)}',
                'pages': 0
            }
    
    @staticmethod
    def _analyze_rtf_file(file_stream, filename, file_size):
        """Analyze RTF file (basic text extraction)"""
        try:
            file_stream.seek(0)
            content = file_stream.read().decode('utf-8', errors='ignore')
            
            # Very basic RTF parsing - extract readable text
            import re
            # Remove RTF control codes
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            text = text.strip()
            
            word_count = len(text.split())
            pages = max(1, word_count // FileAnalyzer.WORDS_PER_PAGE)
            
            return {
                'success': True,
                'content': text,
                'word_count': word_count,
                'pages': pages,
                'file_type': 'rtf',
                'file_size': file_size,
                'filename': filename
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to process RTF file: {str(e)}',
                'pages': 0
            }
    
    @staticmethod
    def get_supported_formats():
        """Get list of supported file formats"""
        return FileAnalyzer.SUPPORTED_FORMATS
    
    @staticmethod
    def format_file_size(size_bytes):
        """Format file size for display"""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024**2:
            return f"{size_bytes/1024:.1f} KB"
        elif size_bytes < 1024**3:
            return f"{size_bytes/(1024**2):.1f} MB"
        else:
            return f"{size_bytes/(1024**3):.1f} GB"

# Global instance
file_analyzer = FileAnalyzer()

def analyze_uploaded_file(file_stream, filename, file_size):
    """Convenience function for file analysis"""
    return file_analyzer.analyze_file(file_stream, filename, file_size)