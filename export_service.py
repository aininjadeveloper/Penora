import os
import logging
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import make_response
import tempfile

class ExportService:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def export_content(self, content, format, title="Generated Content"):
        """Export content in specified format"""
        if format == 'docx':
            return self.create_doc_file(title, content)
        elif format == 'txt':
            return self.create_txt_file(title, content)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def create_doc_file(self, title, content, chapters=None):
        """Create a .docx file from story content"""
        try:
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("Created with Penora AI")
            doc.add_paragraph().add_run().add_break()
            
            if chapters:
                # Multi-chapter story
                for i, chapter in enumerate(chapters, 1):
                    doc.add_heading(f"Chapter {i}", level=1)
                    doc.add_paragraph(chapter)
                    doc.add_paragraph().add_run().add_break()
            else:
                # Single content
                doc.add_paragraph(content)
            
            # Save to bytes
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return {
                'success': True,
                'data': doc_io.getvalue(),
                'filename': f"{title.replace(' ', '_')}.docx"
            }
        except Exception as e:
            logging.error(f"Error creating DOC file: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def create_txt_file(self, title, content, chapters=None):
        """Create a plain text file from story content"""
        try:
            text_content = f"{title}\n"
            text_content += "=" * len(title) + "\n\n"
            text_content += f"Generated on: {datetime.now().strftime('%B %d, %Y')}\n"
            text_content += "Created with Penora AI\n\n"
            
            if chapters:
                # Multi-chapter story
                for i, chapter in enumerate(chapters, 1):
                    text_content += f"Chapter {i}\n"
                    text_content += "-" * 20 + "\n\n"
                    text_content += chapter + "\n\n"
            else:
                # Single content
                text_content += content
            
            return {
                'success': True,
                'data': text_content.encode('utf-8'),
                'filename': f"{title.replace(' ', '_')}.txt"
            }
        except Exception as e:
            logging.error(f"Error creating TXT file: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def create_pdf_file(self, title, content, chapters=None):
        """Create a PDF file from story content"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            # Create PDF buffer
            pdf_buffer = io.BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, 
                                  rightMargin=72, leftMargin=72, 
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leading=14
            )
            
            chapter_style = ParagraphStyle(
                'ChapterTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
                spaceBefore=20
            )
            
            # Build PDF content
            story = []
            
            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Add metadata
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", content_style))
            story.append(Paragraph("Created with Penora AI", content_style))
            story.append(Spacer(1, 20))
            
            if chapters:
                # Multi-chapter story
                for i, chapter in enumerate(chapters, 1):
                    story.append(Paragraph(f"Chapter {i}", chapter_style))
                    story.append(Paragraph(chapter.replace('\n', '<br/>'), content_style))
                    story.append(Spacer(1, 20))
            else:
                # Single content
                story.append(Paragraph(content.replace('\n', '<br/>'), content_style))
            
            # Build PDF
            doc.build(story)
            pdf_buffer.seek(0)
            
            return {
                'success': True,
                'data': pdf_buffer.getvalue(),
                'filename': f"{title.replace(' ', '_')}.pdf"
            }
        except Exception as e:
            logging.error(f"Error creating PDF file: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def get_google_drive_auth_url(self):
        """Get Google Drive authorization URL"""
        # This would require Google Drive API setup
        # For now, return placeholder
        return {
            'success': False,
            'error': 'Google Drive integration requires additional setup'
        }
    
    def get_dropbox_auth_url(self):
        """Get Dropbox authorization URL"""
        # This would require Dropbox API setup
        # For now, return placeholder
        return {
            'success': False,
            'error': 'Dropbox integration requires additional setup'
        }
    
    def prepare_file_response(self, file_data, filename, content_type):
        """Prepare Flask response for file download"""
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

# Initialize export service
export_service = ExportService()